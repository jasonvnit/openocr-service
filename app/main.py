import csv
import gc
import io
import json
import logging
import os
import shutil
import subprocess
import threading
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, unquote, urlparse
from urllib.request import Request, urlopen

from fastapi import Depends, FastAPI, File, Header, HTTPException, Query, UploadFile
from fastapi.concurrency import run_in_threadpool

try:
    from openocr import OpenOCR
except ImportError as exc:  # pragma: no cover - startup guard for misconfigured images
    OpenOCR = None
    OPENOCR_IMPORT_ERROR = exc
else:
    OPENOCR_IMPORT_ERROR = None

logging.basicConfig(
    level=logging.INFO,
    format="[%(asctime)s] %(name)s %(levelname)s: %(message)s",
    datefmt="%Y/%m/%d %H:%M:%S",
)
logger = logging.getLogger(__name__)

TMP_ROOT = Path(os.getenv("OPENOCR_TMP_DIR", "/tmp"))
JOB_ROOT = Path(os.getenv("OPENOCR_JOB_DIR", "/tmp/openocr-jobs"))
API_KEY = os.getenv("API_KEY")
MAX_UPLOAD_BYTES = int(os.getenv("OPENOCR_MAX_UPLOAD_BYTES", str(25 * 1024 * 1024)))
UPLOAD_CHUNK_BYTES = 1024 * 1024
DOC_MAX_PARALLEL_BLOCKS = int(os.getenv("OPENOCR_DOC_MAX_PARALLEL_BLOCKS", "1"))
FILE_URL_TIMEOUT_SECONDS = float(os.getenv("OPENOCR_FILE_URL_TIMEOUT_SECONDS", "30"))
JOB_RESULT_TTL_SECONDS = int(os.getenv("OPENOCR_JOB_RESULT_TTL_SECONDS", str(24 * 60 * 60)))
STT_MODEL_SIZE = os.getenv("OPENOCR_STT_MODEL_SIZE", "small")
STT_COMPUTE_TYPE = os.getenv("OPENOCR_STT_COMPUTE_TYPE", "int8")
STT_LANGUAGE = os.getenv("OPENOCR_STT_LANGUAGE") or None
STT_BEAM_SIZE = int(os.getenv("OPENOCR_STT_BEAM_SIZE", "1"))
STT_VAD_FILTER = os.getenv("OPENOCR_STT_VAD_FILTER", "true").lower() in {"1", "true", "yes", "on"}
OCR_TASKS = {"ocr", "doc"}
IMAGE_SUFFIXES = {".bmp", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}
PDF_SUFFIXES = {".pdf"}
WORD_SUFFIXES = {".doc", ".docx"}
CSV_SUFFIXES = {".csv"}
EXCEL_SUFFIXES = {".xlsm", ".xlsx"}
AUDIO_SUFFIXES = {".aac", ".flac", ".m4a", ".mp3", ".ogg", ".wav", ".webm"}
VIDEO_SUFFIXES = {".avi", ".mkv", ".mov", ".mp4", ".webm"}

_engine_lock = threading.Lock()
_active_engine: Any | None = None
_active_engine_key: str | None = None
_speech_lock = threading.Lock()
_speech_model: Any | None = None
_jobs_lock = threading.Lock()
_jobs: dict[str, dict[str, Any]] = {}
_cleanup_stop_event = threading.Event()


@asynccontextmanager
async def lifespan(app: FastAPI):
    await run_in_threadpool(_warm_up_models)
    JOB_ROOT.mkdir(parents=True, exist_ok=True)
    _cleanup_stop_event.clear()
    cleanup_thread = threading.Thread(target=_cleanup_expired_jobs_loop, daemon=True)
    cleanup_thread.start()
    yield
    _cleanup_stop_event.set()
    cleanup_thread.join(timeout=5)


app = FastAPI(title="openocr-service", lifespan=lifespan)


def _require_api_key(x_api_key: str | None = Header(default=None)) -> None:
    if not API_KEY:
        return

    if x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/extract/file")
async def extract_file(
    file: UploadFile = File(...),
    _: None = Depends(_require_api_key),
) -> dict[str, Any]:
    return await _create_file_job(file)


@app.post("/extract/url")
async def extract_url(
    url: str | None = Query(default=None),
    _: None = Depends(_require_api_key),
) -> dict[str, Any]:
    return await _create_url_job(url)


@app.get("/jobs/{job_id}")
def get_job(
    job_id: str,
    _: None = Depends(_require_api_key),
) -> dict[str, Any]:
    _cleanup_expired_jobs()

    with _jobs_lock:
        job = _jobs.get(job_id)
        if job is None:
            raise HTTPException(status_code=404, detail="Job not found")

        return _public_job(job)


async def _create_file_job(file: UploadFile) -> dict[str, Any]:
    task = _detect_task_from_filename(file.filename)
    _ensure_task_available(task)

    job_id = str(uuid.uuid4())
    work_dir = JOB_ROOT / job_id
    output_dir = work_dir / "output"
    work_dir.mkdir(parents=True, exist_ok=False)
    output_dir.mkdir(parents=True, exist_ok=True)

    source_name = file.filename
    input_path = work_dir / _safe_input_name(source_name, job_id)

    try:
        await run_in_threadpool(_copy_upload, file.file, input_path)
        return _start_job(job_id, task, "file", source_name, None, work_dir, input_path, output_dir)
    except HTTPException:
        shutil.rmtree(work_dir, ignore_errors=True)
        raise
    except Exception as exc:
        shutil.rmtree(work_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        await file.close()


async def _create_url_job(file_url: str | None) -> dict[str, Any]:
    file_url = file_url.strip() if file_url else None
    if not file_url:
        raise HTTPException(status_code=400, detail="url is required.")

    _validate_file_url(file_url)
    task = _detect_task_from_url(file_url)
    _ensure_task_available(task)

    job_id = str(uuid.uuid4())
    work_dir = JOB_ROOT / job_id
    output_dir = work_dir / "output"
    work_dir.mkdir(parents=True, exist_ok=False)
    output_dir.mkdir(parents=True, exist_ok=True)

    source = "youtubeUrl" if task == "youtube" else "fileUrl"
    source_name = _youtube_video_id(file_url) if task == "youtube" else _filename_from_url(file_url)
    input_path = work_dir / _safe_input_name(source_name, job_id)

    try:
        return _start_job(job_id, task, source, source_name, file_url, work_dir, input_path, output_dir)
    except HTTPException:
        shutil.rmtree(work_dir, ignore_errors=True)
        raise
    except Exception as exc:
        shutil.rmtree(work_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=str(exc)) from exc


def _start_job(
    job_id: str,
    task: str,
    source: str,
    source_name: str | None,
    file_url: str | None,
    work_dir: Path,
    input_path: Path,
    output_dir: Path,
) -> dict[str, Any]:
    job = _new_job(job_id, task, source, source_name, file_url, work_dir)
    with _jobs_lock:
        _jobs[job_id] = job

    _log_job_progress(job_id, task, "queued")
    response = _public_job(job)
    thread = threading.Thread(
        target=_run_job,
        args=(job_id, task, input_path, output_dir, file_url),
        daemon=True,
    )
    thread.start()
    return response


def _copy_upload(src: Any, dst_path: Path) -> None:
    total_bytes = 0

    with dst_path.open("wb") as dst:
        while chunk := src.read(UPLOAD_CHUNK_BYTES):
            total_bytes += len(chunk)
            if total_bytes > MAX_UPLOAD_BYTES:
                raise HTTPException(
                    status_code=413,
                    detail=f"Upload too large. Max size is {MAX_UPLOAD_BYTES} bytes.",
                )
            dst.write(chunk)


def _copy_file_url(file_url: str, dst_path: Path) -> None:
    _validate_file_url(file_url)
    request = Request(file_url, headers={"User-Agent": "openocr-service/1.0"})

    try:
        with urlopen(request, timeout=FILE_URL_TIMEOUT_SECONDS) as response:
            content_length = response.headers.get("Content-Length")
            try:
                content_length_bytes = int(content_length) if content_length else 0
            except ValueError:
                content_length_bytes = 0

            if content_length_bytes > MAX_UPLOAD_BYTES:
                raise HTTPException(
                    status_code=413,
                    detail=f"url too large. Max size is {MAX_UPLOAD_BYTES} bytes.",
                )

            _copy_upload(response, dst_path)
    except HTTPException:
        raise
    except HTTPError as exc:
        raise HTTPException(
            status_code=400,
            detail=f"url download failed with HTTP {exc.code}.",
        ) from exc
    except (TimeoutError, URLError, OSError) as exc:
        raise HTTPException(status_code=400, detail=f"url download failed: {exc}") from exc


def _validate_file_url(file_url: str) -> None:
    parsed_url = urlparse(file_url)
    if parsed_url.scheme not in {"http", "https"}:
        raise HTTPException(status_code=400, detail="url must use http or https.")


def _detect_task_from_url(file_url: str) -> str:
    if _youtube_video_id(file_url):
        return "youtube"

    return _detect_task_from_filename(_filename_from_url(file_url))


def _detect_task_from_filename(filename: str | None) -> str:
    suffix = Path(filename or "").suffix.lower()

    if suffix in IMAGE_SUFFIXES:
        return "ocr"

    if suffix in PDF_SUFFIXES:
        return "doc"

    if suffix in WORD_SUFFIXES:
        return "word"

    if suffix in CSV_SUFFIXES:
        return "csv"

    if suffix in EXCEL_SUFFIXES:
        return "excel"

    if suffix in AUDIO_SUFFIXES or suffix in VIDEO_SUFFIXES:
        return "speech"

    raise HTTPException(
        status_code=400,
        detail=(
            "Unsupported file type. Supported types: images, pdf, doc, docx, csv, xlsx, xlsm, "
            "audio, video, and YouTube URLs."
        ),
    )


def _filename_from_url(file_url: str | None) -> str | None:
    if not file_url:
        return None

    parsed_url = urlparse(file_url)
    filename = Path(unquote(parsed_url.path)).name
    return filename or None


def _safe_input_name(filename: str | None, request_id: str) -> str:
    suffix = Path(filename or "").suffix.lower()
    return f"{request_id}{suffix}"


def _new_job(
    job_id: str,
    task: str,
    source: str,
    filename: str | None,
    file_url: str | None,
    work_dir: Path,
) -> dict[str, Any]:
    now = _utc_now()
    return {
        "job_id": job_id,
        "status": "queued",
        "status_url": f"/jobs/{job_id}",
        "task": task,
        "source": source,
        "filename": filename,
        "file_url": file_url,
        "created_at": _format_dt(now),
        "started_at": None,
        "completed_at": None,
        "expires_at": None,
        "result": None,
        "error": None,
        "_work_dir": str(work_dir),
        "_expires_at": None,
    }


def _run_job(
    job_id: str,
    task: str,
    input_path: Path,
    output_dir: Path,
    file_url: str | None,
) -> None:
    _update_job(job_id, status="running", started_at=_format_dt(_utc_now()))
    _log_job_progress(job_id, task, "running")

    try:
        if file_url and task != "youtube":
            _copy_file_url(file_url, input_path)

        result = _run_task(task, input_path, output_dir, file_url)
        _finish_job(
            job_id,
            status="succeeded",
            result=_normalize_result(task, result),
            error=None,
        )
        _log_job_progress(job_id, task, "succeeded")
    except HTTPException as exc:
        _finish_job(
            job_id,
            status="failed",
            result=None,
            error={"status_code": exc.status_code, "detail": exc.detail},
        )
        _log_job_progress(job_id, task, "failed")
    except Exception as exc:
        _finish_job(
            job_id,
            status="failed",
            result=None,
            error={"status_code": 500, "detail": str(exc)},
        )
        _log_job_progress(job_id, task, "failed")
    finally:
        with _jobs_lock:
            job = _jobs.get(job_id)
            work_dir = Path(job["_work_dir"]) if job else None

        if work_dir is not None:
            shutil.rmtree(work_dir, ignore_errors=True)


def _update_job(job_id: str, **updates: Any) -> None:
    with _jobs_lock:
        job = _jobs.get(job_id)
        if job is not None:
            job.update(updates)


def _finish_job(
    job_id: str,
    status: str,
    result: Any,
    error: dict[str, Any] | None,
) -> None:
    now = _utc_now()
    expires_at = now + timedelta(seconds=JOB_RESULT_TTL_SECONDS)
    _update_job(
        job_id,
        status=status,
        completed_at=_format_dt(now),
        expires_at=_format_dt(expires_at),
        result=result,
        error=error,
        _expires_at=expires_at,
    )


def _public_job(job: dict[str, Any]) -> dict[str, Any]:
    return {
        "job_id": job["job_id"],
        "status": job["status"],
        "status_url": job["status_url"],
        "task": job["task"],
        "source": job["source"],
        "filename": job["filename"],
        "created_at": job["created_at"],
        "started_at": job["started_at"],
        "completed_at": job["completed_at"],
        "expires_at": job["expires_at"],
        "result": job["result"] if job["status"] == "succeeded" else None,
        "error": job["error"],
    }


def _cleanup_expired_jobs_loop() -> None:
    while not _cleanup_stop_event.wait(300):
        _cleanup_expired_jobs()


def _cleanup_expired_jobs() -> None:
    now = _utc_now()
    expired_jobs: list[dict[str, Any]] = []

    with _jobs_lock:
        for job_id, job in list(_jobs.items()):
            expires_at = job.get("_expires_at")
            if expires_at is not None and expires_at <= now:
                expired_jobs.append(_jobs.pop(job_id))

    for job in expired_jobs:
        shutil.rmtree(job["_work_dir"], ignore_errors=True)
        _log_job_progress(job["job_id"], job["task"], "expired")


def _log_job_progress(job_id: str, task: str, status: str) -> None:
    logger.info("job_id=%s task=%s status=%s", job_id, task, status)


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _format_dt(value: datetime) -> str:
    return value.isoformat().replace("+00:00", "Z")


def _ensure_task_available(task: str) -> None:
    if task in OCR_TASKS and OPENOCR_IMPORT_ERROR is not None:
        raise HTTPException(
            status_code=500,
            detail=f"openocr-python is not installed correctly: {OPENOCR_IMPORT_ERROR}",
        )


def _run_task(
    task: str,
    input_path: Path,
    output_dir: Path,
    file_url: str | None,
) -> Any:
    if task in OCR_TASKS:
        return _run_openocr(task, input_path, output_dir)

    if task == "word":
        return _run_word_task(input_path)

    if task == "csv":
        return _read_csv_file(input_path)

    if task == "excel":
        return _read_excel_file(input_path)

    if task == "youtube":
        return _read_youtube_transcript(file_url)

    if task == "speech":
        return _run_speech_task(input_path)

    raise ValueError(f"Unsupported task: {task}")


def _run_word_task(input_path: Path) -> Any:
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        return _read_docx_file(input_path)

    if suffix == ".doc":
        return _read_doc_file(input_path)

    raise HTTPException(status_code=400, detail="Unsupported Word file type.")


def _read_docx_file(input_path: Path) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="python-docx is not installed.") from exc

    document = Document(input_path)
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]
    tables = [_docx_table_to_rows(table) for table in document.tables]

    markdown_sections = [*paragraphs]
    for index, table_rows in enumerate(tables, start=1):
        table_markdown = _table_rows_to_markdown(table_rows)
        if table_markdown:
            markdown_sections.append(f"## Table {index}\n\n{table_markdown}")

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "markdown": "\n\n".join(markdown_sections).strip(),
    }


def _read_doc_file(input_path: Path) -> dict[str, Any]:
    antiword_path = shutil.which("antiword")
    if antiword_path is None:
        raise HTTPException(
            status_code=500,
            detail="antiword is not installed; legacy .doc extraction is unavailable.",
        )

    result = subprocess.run(
        [antiword_path, str(input_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        detail = result.stderr.strip() or "antiword failed to extract .doc content."
        raise HTTPException(status_code=500, detail=detail)

    text = result.stdout.strip()
    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n") if paragraph.strip()]
    return {
        "paragraphs": paragraphs,
        "paragraph_count": len(paragraphs),
        "markdown": "\n\n".join(paragraphs).strip(),
    }


def _docx_table_to_rows(table: Any) -> list[list[str]]:
    return [
        [_stringify_cell(cell.text).strip() for cell in row.cells]
        for row in table.rows
    ]


def _read_csv_file(input_path: Path) -> dict[str, Any]:
    raw_content = input_path.read_bytes()
    text = _decode_text(raw_content)
    sample = text[:4096]

    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(io.StringIO(text), dialect)
    rows = [[_stringify_cell(value) for value in row] for row in reader]
    columns = rows[0] if rows else []
    data_rows = rows[1:] if rows else []

    return {
        "columns": columns,
        "rows": data_rows,
        "row_count": len(data_rows),
        "markdown": _rows_to_markdown(columns, data_rows),
    }


def _read_excel_file(input_path: Path) -> dict[str, Any]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="openpyxl is not installed.") from exc

    workbook = load_workbook(input_path, read_only=True, data_only=True)
    sheets: list[dict[str, Any]] = []
    markdown_sections: list[str] = []

    try:
        for worksheet in workbook.worksheets:
            rows = [
                [_stringify_cell(cell) for cell in row]
                for row in worksheet.iter_rows(values_only=True)
            ]
            rows = _trim_empty_rows(rows)
            columns = rows[0] if rows else []
            data_rows = rows[1:] if rows else []

            sheets.append(
                {
                    "name": worksheet.title,
                    "columns": columns,
                    "rows": data_rows,
                    "row_count": len(data_rows),
                }
            )

            markdown_table = _rows_to_markdown(columns, data_rows)
            if markdown_table:
                markdown_sections.append(f"## {worksheet.title}\n\n{markdown_table}")
    finally:
        workbook.close()

    return {
        "sheets": sheets,
        "sheet_count": len(sheets),
        "markdown": "\n\n".join(markdown_sections).strip(),
    }


def _read_youtube_transcript(youtube_url: str | None) -> dict[str, Any]:
    if not youtube_url:
        raise HTTPException(status_code=400, detail="url is required.")

    try:
        from youtube_transcript_api import YouTubeTranscriptApi
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="youtube-transcript-api is not installed.",
        ) from exc

    video_id = _youtube_video_id(youtube_url)
    if not video_id:
        raise HTTPException(status_code=400, detail="Invalid YouTube URL.")

    languages = ["vi", "en", "en-US"]

    try:
        if hasattr(YouTubeTranscriptApi, "get_transcript"):
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=languages)
        else:
            transcript = YouTubeTranscriptApi().fetch(video_id, languages=languages).to_raw_data()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"YouTube transcript unavailable: {exc}") from exc

    segments = [
        {
            "text": str(item.get("text", "")).strip(),
            "start": item.get("start"),
            "duration": item.get("duration"),
        }
        for item in transcript
        if str(item.get("text", "")).strip()
    ]
    markdown = "\n\n".join(segment["text"] for segment in segments).strip()

    return {
        "video_id": video_id,
        "segments": segments,
        "segment_count": len(segments),
        "markdown": markdown,
    }

def _run_speech_task(input_path: Path) -> dict[str, Any]:
    with _speech_lock, _engine_lock:
        _release_openocr_engine()
        model = _get_speech_model()
        segments_iter, info = model.transcribe(
            str(input_path),
            language=STT_LANGUAGE,
            beam_size=STT_BEAM_SIZE,
            vad_filter=STT_VAD_FILTER,
        )
        segments = [
            {
                "start": segment.start,
                "end": segment.end,
                "text": segment.text.strip(),
            }
            for segment in segments_iter
            if segment.text and segment.text.strip()
        ]

    text = " ".join(segment["text"] for segment in segments).strip()
    markdown = "\n\n".join(segment["text"] for segment in segments).strip()

    return {
        "text": text,
        "segments": segments,
        "segment_count": len(segments),
        "language": getattr(info, "language", STT_LANGUAGE),
        "duration": getattr(info, "duration", None),
        "markdown": markdown,
    }


def _get_speech_model() -> Any:
    global _speech_model

    if _speech_model is not None:
        return _speech_model

    try:
        from faster_whisper import WhisperModel
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="faster-whisper is not installed.") from exc

    _speech_model = WhisperModel(
        STT_MODEL_SIZE,
        device="cpu",
        compute_type=STT_COMPUTE_TYPE,
    )
    return _speech_model


def _youtube_video_id(youtube_url: str | None) -> str | None:
    if not youtube_url:
        return None

    parsed_url = urlparse(youtube_url)
    hostname = (parsed_url.hostname or "").lower()

    if hostname in {"youtu.be", "www.youtu.be"}:
        return parsed_url.path.strip("/") or None

    if hostname.endswith("youtube.com"):
        query = parse_qs(parsed_url.query)
        if query.get("v"):
            return query["v"][0]

        path_parts = [part for part in parsed_url.path.split("/") if part]
        if len(path_parts) >= 2 and path_parts[0] in {"embed", "shorts", "live"}:
            return path_parts[1]

    return None


def _decode_text(raw_content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            pass

    return raw_content.decode("utf-8", errors="replace")


def _trim_empty_rows(rows: list[list[str]]) -> list[list[str]]:
    return [row for row in rows if any(value for value in row)]


def _stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _rows_to_markdown(columns: list[str], rows: list[list[str]]) -> str:
    if not columns and not rows:
        return ""

    column_count = max([len(columns), *(len(row) for row in rows)] or [0])
    if column_count == 0:
        return ""

    headers = _pad_row(columns, column_count)
    if not any(headers):
        headers = [f"Column {index + 1}" for index in range(column_count)]

    table_lines = [
        "| " + " | ".join(_escape_markdown_cell(value) for value in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]

    for row in rows:
        values = _pad_row(row, column_count)
        table_lines.append("| " + " | ".join(_escape_markdown_cell(value) for value in values) + " |")

    return "\n".join(table_lines)


def _table_rows_to_markdown(rows: list[list[str]]) -> str:
    rows = _trim_empty_rows(rows)
    if not rows:
        return ""

    columns = rows[0]
    data_rows = rows[1:]
    return _rows_to_markdown(columns, data_rows)


def _pad_row(row: list[str], column_count: int) -> list[str]:
    return [*row, *([""] * max(column_count - len(row), 0))][:column_count]


def _escape_markdown_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def _run_openocr(task: str, input_path: Path, output_dir: Path) -> Any:
    with _engine_lock:
        engine = _get_engine(task)

        if task == "ocr":
            return engine(str(input_path), save_dir=str(output_dir), is_visualize=False)

        if task == "doc":
            return engine(str(input_path))

        raise ValueError(f"Unsupported task: {task}")


def _normalize_result(task: str, result: Any) -> Any:
    if task != "ocr":
        normalized = _to_jsonable(result)
        return _with_markdown(task, normalized)

    if result is None:
        return _with_markdown(task, {"items": [], "timings": []})

    if (
        isinstance(result, tuple)
        and len(result) == 2
    ):
        lines, timings = result
        if lines is None and timings is None:
            return _with_markdown(task, {"items": [], "timings": []})

        if not lines:
            return _with_markdown(task, {"items": [], "timings": _to_jsonable(timings) or []})

        normalized = {
            "items": [_parse_ocr_line(line) for line in lines],
            "timings": _to_jsonable(timings) or [],
        }
        return _with_markdown(task, normalized)

    normalized = _to_jsonable(result)
    return _with_markdown(task, normalized)


def _with_markdown(task: str, result: Any) -> dict[str, Any]:
    markdown = _result_to_markdown(task, result)

    if isinstance(result, dict):
        return {**result, "markdown": markdown}

    return {"data": result, "markdown": markdown}


def _result_to_markdown(task: str, result: Any) -> str:
    if task == "ocr":
        return _ocr_result_to_markdown(result)

    markdown = _find_markdown(result)
    if markdown:
        return markdown

    return "```json\n" + json.dumps(result, ensure_ascii=False, indent=2) + "\n```"


def _ocr_result_to_markdown(result: Any) -> str:
    if not isinstance(result, dict):
        return ""

    sections: list[str] = []
    for item in result.get("items", []):
        if not isinstance(item, dict):
            continue

        image_name = item.get("image_name")
        if image_name:
            sections.append(f"## {image_name}")

        texts = _extract_ocr_texts(item.get("ocr"))
        sections.extend(texts)

    return "\n\n".join(section for section in sections if section).strip()


def _extract_ocr_texts(value: Any) -> list[str]:
    texts: list[str] = []

    if isinstance(value, dict):
        for key in ("text", "transcription", "rec_text", "label"):
            text = value.get(key)
            if isinstance(text, str) and text.strip():
                texts.append(text.strip())

        for child in value.values():
            texts.extend(_extract_ocr_texts(child))
        return _dedupe_texts(texts)

    if isinstance(value, (list, tuple)):
        if len(value) >= 2 and isinstance(value[1], (list, tuple)):
            candidate = value[1][0] if value[1] else None
            if isinstance(candidate, str) and candidate.strip():
                texts.append(candidate.strip())

        for child in value:
            texts.extend(_extract_ocr_texts(child))
        return _dedupe_texts(texts)

    return texts


def _find_markdown(value: Any) -> str | None:
    if isinstance(value, dict):
        for key in ("markdown", "md", "markdown_text", "md_content"):
            markdown = value.get(key)
            if isinstance(markdown, str) and markdown.strip():
                return markdown

        for child in value.values():
            markdown = _find_markdown(child)
            if markdown:
                return markdown

    if isinstance(value, list):
        for child in value:
            markdown = _find_markdown(child)
            if markdown:
                return markdown

    return None


def _dedupe_texts(texts: list[str]) -> list[str]:
    seen: set[str] = set()
    deduped: list[str] = []

    for text in texts:
        if text not in seen:
            seen.add(text)
            deduped.append(text)

    return deduped


def _parse_ocr_line(line: Any) -> Any:
    if not isinstance(line, str) or "\t" not in line:
        return _to_jsonable(line)

    image_name, raw_result = line.split("\t", 1)
    return {
        "image_name": image_name,
        "ocr": _parse_json_string(raw_result),
    }


def _get_engine(task: str) -> Any:
    global _active_engine, _active_engine_key

    if _active_engine is not None and _active_engine_key == task:
        return _active_engine

    _release_openocr_engine()
    _release_speech_model()

    if task == "ocr":
        # infer_e2e.py compares use_gpu with `==` against the string 'false'.
        _active_engine = OpenOCR(task="ocr", mode="mobile", use_gpu="false")
    elif task == "doc":
        # infer_doc_onnx.py / infer_unirec_onnx.py compare use_gpu with `is False`,
        # so a string here silently falls through to GPU auto-detection instead of
        # forcing CPU.
        _active_engine = OpenOCR(
            task="doc",
            use_layout_detection=True,
            use_gpu=False,
            max_parallel_blocks=DOC_MAX_PARALLEL_BLOCKS,
        )
    else:
        raise ValueError(f"Unsupported task: {task}")

    _active_engine_key = task
    return _active_engine


def _release_openocr_engine() -> None:
    global _active_engine, _active_engine_key

    _active_engine = None
    _active_engine_key = None
    gc.collect()


def _release_speech_model() -> None:
    global _speech_model

    _speech_model = None
    gc.collect()


def _warm_up_models() -> None:
    if OPENOCR_IMPORT_ERROR is not None:
        logger.warning("Skipping OpenOCR warm-up: %s", OPENOCR_IMPORT_ERROR)
        return

    for task in ("ocr", "doc"):
        try:
            _get_engine(task)
            logger.info("OpenOCR %s model warm-up completed", task)
        except Exception as exc:
            logger.warning("OpenOCR %s model warm-up failed: %s", task, exc)


def _to_jsonable(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return _parse_json_string(value) if isinstance(value, str) else value

    if isinstance(value, dict):
        return {str(k): _to_jsonable(v) for k, v in value.items()}

    if isinstance(value, (list, tuple, set)):
        return [_to_jsonable(item) for item in value]

    if hasattr(value, "tolist"):
        return _to_jsonable(value.tolist())

    if hasattr(value, "item"):
        try:
            return value.item()
        except Exception:
            pass

    return str(value)


def _parse_json_string(value: str) -> Any:
    stripped = value.strip()
    if not stripped or stripped[0] not in "[{":
        return value

    try:
        return json.loads(stripped)
    except json.JSONDecodeError:
        return value
